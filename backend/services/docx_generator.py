"""Professional bilingual Word generator.

Output is a single continuous paper: no per-question boxes, no separate table
for every question, English followed by NCERT-style Hindi, source figures kept
inline, and every question is kept together on one page when Word can fit it.
"""

from pathlib import Path
import re

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TOKEN_RE = re.compile(r"__MATH_(\d+)__")
QUESTION_RE = re.compile(r"^\s*(?:Q\.?\s*)?(\d{1,3})[\.\)]\s*(.*)$", re.S)
SECTION_NAMES = {
    "PHYSICS": "भौतिकी (PHYSICS)",
    "CHEMISTRY": "रसायन विज्ञान (CHEMISTRY)",
    "MATHEMATICS": "गणित (MATHEMATICS)",
    "BIOLOGY": "जीवविज्ञान (BIOLOGY)",
    "BOTANY": "वनस्पति विज्ञान (BOTANY)",
    "ZOOLOGY": "प्राणी विज्ञान (ZOOLOGY)",
}


def _is_hindi(text: str) -> bool:
    return any("\u0900" <= c <= "\u097F" for c in str(text or ""))


def _set_run_font(run, text, size=9.7, bold=False, font=None):
    run.bold = bold
    run.font.name = font or ("Noto Sans Devanagari" if _is_hindi(text) else "Times New Roman")
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), run.font.name)
    rfonts.set(qn("w:hAnsi"), run.font.name)
    rfonts.set(qn("w:eastAsia"), run.font.name)
    rfonts.set(qn("w:cs"), run.font.name)


def _set_cell_margins(cell, top=0, start=70, bottom=0, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def _set_fixed_table(table):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    return p


def _add_plain_paragraph(cell, text, *, size=9.6, bold=False, hindi=False, keep=True, indent=0):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = keep
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    _set_run_font(r, text, size=size, bold=bold, font="Noto Sans Devanagari" if hindi else "Times New Roman")
    return p


def _add_omml_math(parent_paragraph, text, superscript=False, subscript=False):
    """Insert a real Word equation run (OMML), not a screenshot.

    For PDF math spans we preserve the exact extracted symbol text and apply
    superscript/subscript structure when PyMuPDF marks it in the source.
    """
    math_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    def E(tag):
        return OxmlElement("m:" + tag)
    omath = E("oMath")
    mr = E("r")
    mt = E("t")
    mt.text = text
    mr.append(mt)
    if superscript:
        ss = E("sSup")
        e = E("e"); e.append(mr)
        sup = E("sup"); sup_r = E("r"); sup_t = E("t"); sup_t.text = " "
        sup_r.append(sup_t); sup.append(sup_r)
        ss.append(e); ss.append(sup)
        omath.append(ss)
    elif subscript:
        ss = E("sSub")
        e = E("e"); e.append(mr)
        sub = E("sub"); sub_r = E("r"); sub_t = E("t"); sub_t.text = " "
        sub_r.append(sub_t); sub.append(sub_r)
        ss.append(e); ss.append(sub)
        omath.append(ss)
    else:
        omath.append(mr)
    parent_paragraph._p.append(omath)


def _add_rich_line(cell, line):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1.2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    for span in line.get("spans", []):
        text = span.get("text", "")
        if not text:
            continue
        flags = int(span.get("flags", 0) or 0)
        font = str(span.get("font", ""))
        if "math" in font.lower() or (flags & 1):
            _add_omml_math(p, text, superscript=bool(flags & 1), subscript=bool(flags & 32))
            continue
        run = p.add_run(text)
        _set_run_font(run, text, size=min(max(float(span.get("size", 10) or 10), 8), 11))
        run.italic = bool(flags & 2)
        run.bold = bool(flags & 16)
    return p


def _add_text_with_tokens(cell, text, math_values, *, hindi=False, size=9.6):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1.2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    pos = 0
    for match in TOKEN_RE.finditer(str(text or "")):
        if match.start() > pos:
            part = text[pos:match.start()]
            r = p.add_run(part)
            _set_run_font(r, part, size=size, font="Noto Sans Devanagari" if hindi else "Times New Roman")
        idx = int(match.group(1))
        value = math_values[idx] if 0 <= idx < len(math_values) else ""
        if value:
            _add_omml_math(p, value)
        pos = match.end()
    if pos < len(text):
        part = text[pos:]
        r = p.add_run(part)
        _set_run_font(r, part, size=size, font="Noto Sans Devanagari" if hindi else "Times New Roman")
    return p


def _add_picture(cell, image_path, width=2.8):
    if not image_path or not Path(image_path).exists():
        return
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(image_path, width=Inches(width))


def _strip_number(text: str):
    match = QUESTION_RE.match(text or "")
    if not match:
        return "", text
    return match.group(1), match.group(2).strip()


def _add_question_language(cell, blocks, language_key, question_number, hindi=False):
    first_question_block = True
    for block in blocks:
        if block.get("type") == "image":
            _add_picture(cell, block.get("image_path"), min(3.2, max(1.25, float(block.get("width", 250)) / 105)))
            continue

        text = str(block.get(language_key, "") or "").strip()
        if not text:
            continue

        if first_question_block:
            _, text_without_number = _strip_number(text)
            text = text_without_number
            first_question_block = False

        if language_key == "english" and block.get("lines"):
            # Preserve original typography and math spans from the PDF.
            lines = block["lines"]
            for line_index, line in enumerate(lines):
                line_text = "".join(s.get("text", "") for s in line.get("spans", []))
                if line_index == 0:
                    _, line_text_without_number = _strip_number(line_text)
                    if line_text != line_text_without_number:
                        # Rebuild first line after removing question number.
                        rebuilt = []
                        remaining = line_text_without_number
                        for span in line.get("spans", []):
                            stext = span.get("text", "")
                            if not remaining:
                                break
                            if stext and remaining.endswith(stext):
                                rebuilt = [span]
                                break
                        if rebuilt:
                            line = dict(line); line["spans"] = rebuilt
                _add_rich_line(cell, line)
            continue

        _add_text_with_tokens(cell, text, block.get("math_values", []), hindi=hindi)


def _add_question_row(table, question):
    row = table.add_row()
    _prevent_row_split(row)
    cell = row.cells[0]
    _set_cell_margins(cell, top=25, start=60, bottom=35, end=60)
    _clear_cell(cell)

    number = int(question.get("number", 0))
    blocks = question.get("blocks", [])

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_together = True
    r = p.add_run(f"{number}.")
    _set_run_font(r, r.text, size=10.5, bold=True)

    # English first, then Hindi. This is a single continuous question block.
    _add_question_language(cell, blocks, "english", number, hindi=False)

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_together = True
    r = p.add_run("हिंदी:")
    _set_run_font(r, r.text, size=9.2, bold=True, font="Noto Sans Devanagari")

    _add_question_language(cell, blocks, "hindi", number, hindi=True)


def _add_section_heading(doc, section_name):
    title = SECTION_NAMES.get(section_name, section_name)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    _set_run_font(r, title, size=13, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # thin rule below the section title
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8"); bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "808080")
    pbdr.append(bottom); pPr.append(pbdr)


def generate_docx(data: dict, output_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9.6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("I PUC JEE (MAINS)")
    _set_run_font(r, r.text, size=17, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Bilingual Question Paper  •  English + हिंदी")
    _set_run_font(r, r.text, size=10.2, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Instructions: English is followed immediately by its NCERT-style Hindi translation. Mathematical expressions are preserved as Word math where source spans permit.")
    _set_run_font(r, r.text, size=8.2)

    last_section = None
    table = None
    for question in data.get("questions", []):
        section_name = str(question.get("section") or "").upper()
        if section_name != last_section:
            if table is not None:
                # Finish the current continuous question container.
                table = None
            if section_name:
                _add_section_heading(doc, section_name)
            table = doc.add_table(rows=0, cols=1)
            _set_fixed_table(table)
            _remove_table_borders(table)
            table.columns[0].width = Inches(7.15)
            last_section = section_name
        if table is None:
            table = doc.add_table(rows=0, cols=1)
            _set_fixed_table(table); _remove_table_borders(table)
        _add_question_row(table, question)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("I PUC JEE (MAINS)  •  Bilingual Paper  •  Page ")
    _set_run_font(run, run.text, size=8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)

    doc.save(output_path)
